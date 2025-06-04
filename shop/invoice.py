from docx import Document
from io import BytesIO

def generate_invoice_docx(order):
    doc = Document()
    doc.add_heading(f'Invoice for Order #{order.id}', level=1)
    
    doc.add_paragraph(f'Customer: {order.name}')
    doc.add_paragraph(f'Phone: {order.phone}')
    doc.add_paragraph(f'Address: {order.address}')
    doc.add_paragraph(f'Pincode: {order.pincode}')
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Price'
    hdr_cells[3].text = 'Total'

    total_amount = 0
    for item in order.items.all():
        row_cells = table.add_row().cells
        row_cells[0].text = item.product_name
        row_cells[1].text = str(item.quantity)
        row_cells[2].text = f'{item.price:.2f}'
        total_price = item.quantity * item.price
        row_cells[3].text = f'{total_price:.2f}'
        total_amount += total_price

    doc.add_paragraph('')
    doc.add_paragraph(f'Total Amount: {total_amount:.2f}')

    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f
